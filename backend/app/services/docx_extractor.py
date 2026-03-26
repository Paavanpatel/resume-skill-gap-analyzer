"""
DOCX text extraction using python-docx.

DOCX files are ZIP archives containing XML. python-docx parses the XML
structure to extract paragraphs, tables, headers, and footers.

Key design choice: we extract text from paragraphs AND tables. Many resumes
use tables for layout (skills grids, two-column formats). Ignoring tables
would miss a large chunk of content in modern resume templates.
"""

import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


class DOCXExtractionError(Exception):
    """Raised when DOCX text extraction fails."""
    pass


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract all text from a DOCX file, including table content.

    The extraction order follows the document's reading order:
    paragraphs are extracted first, then table cells. This preserves
    the natural flow of a resume (heading -> content -> skills table).

    Args:
        file_content: Raw bytes of the DOCX file.

    Returns:
        The full extracted text as a single string.

    Raises:
        DOCXExtractionError: If the DOCX cannot be parsed or has no content.
    """
    try:
        doc = Document(io.BytesIO(file_content))
    except Exception as e:
        raise DOCXExtractionError(f"Could not open DOCX file: {e}") from e

    text_parts: list[str] = []

    # Extract paragraph text (the main body of the resume)
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    # Extract table text (skills grids, structured sections).
    # Each row becomes a single line with cell values separated by tabs,
    # which makes it easy for the section parser to handle later.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append("\t".join(cells))

    if not text_parts:
        raise DOCXExtractionError(
            "No text could be extracted from the DOCX file. "
            "The document appears to be empty."
        )

    return "\n".join(text_parts)
